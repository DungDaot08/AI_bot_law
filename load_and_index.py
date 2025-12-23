import os
import re
from typing import List
from docx import Document

from langchain.schema import Document as LCDocument
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

# =====================
# CONFIG
# =====================
DATA_DIR = "./data"
CHROMA_DIR = "./chroma_db"
DEBUG_DIR = "./debug_chunks"

COLLECTION_NAME = "legal_docs"
EMBED_MODEL = "bkai-foundation-models/vietnamese-bi-encoder"

os.makedirs(DEBUG_DIR, exist_ok=True)

# =====================
# UTIL
# =====================


def is_temp_file(filename: str) -> bool:
    return filename.startswith("~$")


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_law_name(paragraphs: List[str]) -> str:
    for p in paragraphs[:10]:
        if p.isupper() and len(p) > 10:
            return p
        if p.startswith("LUẬT") or p.startswith("NGHỊ ĐỊNH") or p.startswith("THÔNG TƯ"):
            return p
    return ""


def has_article_structure(paragraphs: List[str]) -> bool:
    for p in paragraphs:
        if re.match(r"^Điều\s+\d+", p):
            return True
    return False

# =====================
# PARSE DOCX
# =====================


def parse_law_docx(path: str) -> List[LCDocument]:
    doc = Document(path)
    raw_paragraphs = [clean_text(p.text)
                      for p in doc.paragraphs if p.text.strip()]
    if not raw_paragraphs:
        return []

    law_name = extract_law_name(raw_paragraphs)
    filename = os.path.basename(path)

    documents = []

    # =====================
    # CASE 1: CÓ ĐIỀU
    # =====================
    if has_article_structure(raw_paragraphs):
        current_article = ""
        current_clause = ""
        buffer = []

        for p in raw_paragraphs:
            article_match = re.match(r"^(Điều\s+\d+\.?.*)", p)
            clause_match = re.match(r"^(\d+\.|\d+\))\s+", p)

            if article_match:
                if buffer:
                    documents.append(
                        LCDocument(
                            page_content="\n".join(buffer),
                            metadata={
                                "law_name": law_name,
                                "article": current_article,
                                "clause": current_clause,
                                "source": filename,
                            },
                        )
                    )
                    buffer = []

                current_article = article_match.group(1)
                current_clause = ""
                buffer.append(p)
                continue

            if clause_match:
                current_clause = clause_match.group(1)
                buffer.append(p)
                continue

            buffer.append(p)

        if buffer:
            documents.append(
                LCDocument(
                    page_content="\n".join(buffer),
                    metadata={
                        "law_name": law_name,
                        "article": current_article,
                        "clause": current_clause,
                        "source": filename,
                    },
                )
            )

    # =====================
    # CASE 2: KHÔNG CÓ ĐIỀU → SEMANTIC CHUNK
    # =====================
    else:
        chunk_size = 4
        for i in range(0, len(raw_paragraphs), chunk_size):
            chunk = raw_paragraphs[i: i + chunk_size]
            documents.append(
                LCDocument(
                    page_content="\n".join(chunk),
                    metadata={
                        "law_name": law_name,
                        "article": "Nội dung hướng dẫn",
                        "clause": "",
                        "source": filename,
                    },
                )
            )

    return documents

# =====================
# DEBUG EXPORT
# =====================


def export_debug_chunks(filename: str, docs: List[LCDocument]):
    debug_path = os.path.join(
        DEBUG_DIR, filename.replace(".docx", ".txt")
    )

    with open(debug_path, "w", encoding="utf-8") as f:
        for i, doc in enumerate(docs, 1):
            f.write(f"\n===== CHUNK {i} =====\n")
            f.write(f"LAW: {doc.metadata.get('law_name')}\n")
            f.write(f"ARTICLE: {doc.metadata.get('article')}\n")
            f.write(f"CLAUSE: {doc.metadata.get('clause')}\n")
            f.write(f"SOURCE: {doc.metadata.get('source')}\n\n")
            f.write(doc.page_content)
            f.write("\n")

    print(f"📝 Đã xuất debug file: {debug_path}")

# =====================
# MAIN
# =====================


def main():
    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)

    vectordb = Chroma(
        persist_directory=CHROMA_DIR,
        collection_name=COLLECTION_NAME,
        embedding_function=embeddings,
    )

    all_docs = []

    for filename in os.listdir(DATA_DIR):
        if not filename.endswith(".docx"):
            continue
        if is_temp_file(filename):
            continue

        path = os.path.join(DATA_DIR, filename)
        print(f"📄 Đang xử lý: {filename}")

        docs = parse_law_docx(path)
        print(f"   → {len(docs)} chunks")

        if docs:
            export_debug_chunks(filename, docs)
            all_docs.extend(docs)

    if not all_docs:
        print("❌ Không có dữ liệu nào để index")
        return

    vectordb.add_documents(all_docs)
    vectordb.persist()

    print(f"\n✅ Đã index {len(all_docs)} chunks vào ChromaDB")


# =====================
if __name__ == "__main__":
    main()
